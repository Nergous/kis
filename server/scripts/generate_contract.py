import sys
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def generate_contract(data, output_path):
    print("----------------CREATING DOCUMENT----------------------")
    doc = Document()
    
    # Add title
    title = doc.add_heading('Договор купли-продажи', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add date
    current_date = datetime.now().strftime("%d.%m.%Y")
    date_para = doc.add_paragraph(f"Дата: {current_date}")
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # Add customer info
    doc.add_heading('1. Информация о заказчике', level=2)
    doc.add_paragraph(f"Адрес доставки: {data['address']}")
    doc.add_paragraph(f"Телефон получателя: {data['recipient_phone']}")
    
    # Add delivery info
    doc.add_heading('2. Условия поставки', level=2)
    delivery_date = datetime.strptime(data['delivery_date'], "%Y-%m-%dT%H:%M:%SZ").strftime("%d.%m.%Y")
    doc.add_paragraph(f"Дата поставки: {delivery_date}")
    doc.add_paragraph(f"Условия оплаты: {translate_payment_terms(data['payment_terms'])}")
    
    # Add order items
    doc.add_heading('3. Состав заказа', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Table headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Товар'
    hdr_cells[2].text = 'Количество'
    hdr_cells[3].text = 'Цена'
    
    # Add items
    for idx, item in enumerate(data['order_content'], start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = item['product_name']
        row_cells[2].text = str(item['quantity'])
        row_cells[3].text = f"{item['price']:.2f} руб."
    
    # Add total
    doc.add_paragraph(f"Итого: {data['total_price']:.2f} руб.", style='Intense Quote')
    
    # Save document
    print("----------------SAVING CONTRACT----------------------")
    doc.save(output_path)
    return output_path

def translate_payment_terms(term):
    terms = {
        'prepayment': 'Предоплата',
        'postpayment': 'Постоплата',
        'full_payment': 'Полная оплата'
    }
    return terms.get(term, term)

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python generate_contract.py <json_data> <output_path>")
        sys.exit(1)

    print("----------------GENERATING CONTRACT----------------------")
    
    json_data = sys.argv[1]
    output_path = sys.argv[2]
    
    try:
        data = json.loads(json_data)
        full_path = generate_contract(data, output_path)
        print(full_path)
    except Exception as e:
        print(f"Error generating contract: {str(e)}", file=sys.stderr)
        sys.exit(1)