import sys
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def generate_receipt(data, output_path):
    print("----------------CREATING RECEIPT----------------------")
    doc = Document()
    
    # Add title
    title = doc.add_heading('Чек об оплате', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add date
    try:
        try:

            payment_date = datetime.strptime(data['payment_time'], "%Y-%m-%dT%H:%M:%SZ")
        except ValueError:
            try:
                payment_date = datetime.strptime(data['payment_time'], "%Y-%m-%dT%H:%M:%S%z")
            except ValueError:
                payment_date = datetime.now()
        date_str = payment_date.strftime("%d.%m.%Y %H:%M")
    except Exception as e:
        print(f"Warning: failed to parse payment time: {str(e)}")
        date_str = datetime.now().strftime("%d.%m.%Y %H:%M")
        
    date_para = doc.add_paragraph(f"Дата оплаты: {payment_date}")
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # Add company info
    doc.add_heading('Реквизиты компании:', level=2)
    doc.add_paragraph("ООО 'Лесопилька'")
    doc.add_paragraph("ИНН: 1234567890")
    doc.add_paragraph("КПП: 987654321")
    doc.add_paragraph("БИК: 044525225")
    doc.add_paragraph("Р/с: 40702810500000012345")
    doc.add_paragraph("К/с: 30101810400000000225")
    doc.add_paragraph("Банк: ПАО 'Лесоповал'")
    doc.add_paragraph("Адрес: г. Лесов, ул. Лесная, д. 1")
    
    # Add customer info
    doc.add_heading('Плательщик:', level=2)
    customer = data['customer']
    if customer['customer_type'] == 'juri':
        doc.add_paragraph(f"Организация: {customer['name']}")
        doc.add_paragraph(f"ИНН: {customer['inn']}")
    else:
        doc.add_paragraph(f"ФИО: {customer['surname']} {customer['first_name']} {customer['patronymic']}")
    doc.add_paragraph(f"Email: {customer['email']}")
    
    # Add payment info
    doc.add_heading('Информация об оплате:', level=2)
    doc.add_paragraph(f"Номер заказа: {data['order_id']}")
    doc.add_paragraph(f"Сумма оплаты: {data['amount']:.2f} руб.")
    
    payment_terms_translation = {
        'prepayment': 'Предоплата',
        'postpayment': 'Постоплата',
        'full_payment': 'Полная оплата'
    }
    doc.add_paragraph(f"Способ расчета: {payment_terms_translation.get(data['payment_terms'], data['payment_terms'])}")
    
    # Add order items
    doc.add_heading('Состав заказа:', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Table headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Товар'
    hdr_cells[2].text = 'Количество'
    hdr_cells[3].text = 'Цена'
    
    # Add items
    for idx, item in enumerate(data['order_content'], start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = item['product']['name']
        row_cells[2].text = str(item['quantity'])
        row_cells[3].text = f"{item['price']:.2f} руб."
    
    # Add total
    doc.add_paragraph(f"Итого: {data['total_price']:.2f} руб.", style='Intense Quote')
    
    # Add signature
    doc.add_paragraph("\n\n")
    doc.add_paragraph("_________________________          _________________________")
    doc.add_paragraph("Подпись плательщика                     Подпись сотрудника")
    
    # Save document
    print("----------------SAVING RECEIPT----------------------")
    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    print("AAAAAAAAAAAAAAAAAAAA ПЕНИСЫ")
    if len(sys.argv) != 3:
        print("Usage: python generate_receipt.py <json_data> <output_path>")
        sys.exit(1)

    print("----------------GENERATING RECEIPT----------------------")
    
    json_data = sys.argv[1]
    output_path = sys.argv[2]
    
    try:
        data = json.loads(json_data)
        full_path = generate_receipt(data, output_path)
        print(full_path)
    except Exception as e:
        print(f"Error generating receipt: {str(e)}", file=sys.stderr)
        sys.exit(1)