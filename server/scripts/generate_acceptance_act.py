import sys
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def translate_payment_terms(term):
    terms = {
        'prepayment': 'Предоплата',
        'postpayment': 'Постоплата',
        'full_payment': 'Полная оплата'
    }
    return terms.get(term, term)

def generate_acceptance_act(data, output_path):
    print("----------------CREATING ACCEPTANCE ACT----------------------")
    doc = Document()
    
    # Add title
    title = doc.add_heading('Акт получения заказа', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add document number and date
    current_date = datetime.now().strftime("%d.%m.%Y")
    doc.add_paragraph(f"Акт № {data['order_id']} от {current_date}")
    
    # Add company info
    doc.add_heading('1. Реквизиты поставщика:', level=2)
    doc.add_paragraph("ООО 'Лесопилька'")
    doc.add_paragraph("ИНН: 1234567890")
    doc.add_paragraph("КПП: 987654321")
    doc.add_paragraph("Юридический адрес: г. Москва, ул. Ленина, д. 1")
    doc.add_paragraph("Телефон: +7 (495) 123-45-67")
    
    # Add customer info
    doc.add_heading('2. Реквизиты получателя:', level=2)
    customer = data['customer']
    if customer['customer_type'] == 'juri':
        doc.add_paragraph(f"Организация: {customer['name']}")
        doc.add_paragraph(f"ИНН: {customer['inn']}")
    else:
        doc.add_paragraph(f"ФИО: {customer['surname']} {customer['first_name']} {customer['patronymic']}")
    doc.add_paragraph(f"Адрес: {data['address']}")
    doc.add_paragraph(f"Телефон: {data['recipient_phone']}")
    
    # Add order info
    doc.add_heading('3. Информация о заказе:', level=2)
    doc.add_paragraph(f"Номер заказа: {data['order_id']}")
    
    try:
        delivery_date = datetime.strptime(data['delivery_date'], "%Y-%m-%dT%H:%M:%SZ").strftime("%d.%m.%Y")
    except:
        delivery_date = datetime.now().strftime("%d.%m.%Y")
    
    doc.add_paragraph(f"Дата получения: {delivery_date}")
    doc.add_paragraph(f"Способ расчета: {translate_payment_terms(data['payment_terms'])}")
    
    # Add order items
    doc.add_heading('4. Полученные товары:', level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Table headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Наименование'
    hdr_cells[2].text = 'Количество'
    hdr_cells[3].text = 'Цена'
    hdr_cells[4].text = 'Сумма'
    
    # Add items
    for idx, item in enumerate(data['order_content'], start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = item['product_name']
        row_cells[2].text = str(item['quantity'])
        row_cells[3].text = f"{item['price']:.2f} руб."
        row_cells[4].text = f"{item['total_product_price']:.2f} руб."
    
    # Add total
    doc.add_paragraph(f"Итого: {data['total_price']:.2f} руб.", style='Intense Quote')
    
    # Add acceptance clause
    doc.add_paragraph("\nТовар получен в полном объеме, претензий по количеству и качеству не имею.")
    
    # Add signatures
    doc.add_paragraph("\n\n_________________________          _________________________")
    doc.add_paragraph("Подпись получателя                     Подпись поставщика")
    doc.add_paragraph(f"Дата: {current_date}")
    
    # Save document
    print("----------------SAVING ACCEPTANCE ACT----------------------")
    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python generate_acceptance_act.py <json_data> <output_path>")
        sys.exit(1)

    print("----------------GENERATING ACCEPTANCE ACT----------------------")
    
    json_data = sys.argv[1]
    output_path = sys.argv[2]
    
    try:
        data = json.loads(json_data)
        print(data)
        full_path = generate_acceptance_act(data, output_path)
        print(full_path)
    except Exception as e:
        print(f"Error generating acceptance act: {str(e)}", file=sys.stderr)
        sys.exit(1)