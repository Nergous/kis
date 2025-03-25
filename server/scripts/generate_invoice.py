import sys
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def generate_invoice(data, output_path):
    print("----------------CREATING INVOICE DOCUMENT----------------------")
    doc = Document()
    
    # Add title
    title = doc.add_heading('Товарно-Транспортная Накладная', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add document number and date
    current_date = datetime.now().strftime("%d.%m.%Y")
    doc.add_paragraph(f"Номер заказа: {data.get('order_id', 'N/A')}")
    date_para = doc.add_paragraph(f"Дата составления: {current_date}")
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # Add goods table
    doc.add_heading('Товары', level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Table headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Наименование товара'
    hdr_cells[2].text = 'Кол-во'
    hdr_cells[3].text = 'Цена за ед.'
    hdr_cells[4].text = 'Сумма'
    
    # Add items
    for idx, item in enumerate(data['order_content'], start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = item['product_name']
        row_cells[2].text = str(item['quantity'])
        row_cells[3].text = f"{item['price']:.2f} руб."
        row_cells[4].text = f"{item['total_product_price']:.2f} руб."
    
    # Add totals
    doc.add_paragraph(f"Итого товаров: {len(data['order_content'])} наименований")
    doc.add_paragraph(f"Общая сумма: {data['total_price']:.2f} руб.", style='Intense Quote')
    
    # Add signatures
    doc.add_heading('3. Подписи сторон', level=2)
    doc.add_paragraph("Отправитель: _________________________")
    
    # Save document
    print("----------------SAVING INVOICE----------------------")
    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python generate_invoice.py <json_data> <output_path>")
        sys.exit(1)

    print("----------------GENERATING INVOICE----------------------")
    
    json_data = sys.argv[1]
    output_path = sys.argv[2]
    
    try:
        data = json.loads(json_data)
        full_path = generate_invoice(data, output_path)
        print(full_path)
    except Exception as e:
        print(f"Error generating invoice: {str(e)}", file=sys.stderr)
        sys.exit(1)